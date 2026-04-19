"""
Document processor for MiniChat application.
Handles document parsing, text extraction, and chunking for RAG.
"""
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
import os

from miniLM.src.database.vector_db import DocumentChunk
from miniLM.src.utils.helpers import generate_uuid, sanitize_filename
from miniLM.src.utils.logger import get_logger
from miniLM.config.settings import get_settings

logger = get_logger(__name__)

# Supported file extensions
SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md'}


class DocumentProcessingError(Exception):
    """Exception raised when document processing fails."""
    def __init__(self, message: str, user_message: str, recoverable: bool = True):
        super().__init__(message)
        self.user_message = user_message
        self.recoverable = recoverable


@dataclass
class ProcessedDocument:
    """Represents a processed document with its chunks."""
    source: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]


class DocumentProcessor:
    """Processes documents for RAG pipeline."""
    
    def __init__(
        self,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None
    ):
        """
        Initialize DocumentProcessor.
        
        Args:
            chunk_size: Size of text chunks. Defaults to settings value.
            chunk_overlap: Overlap between chunks. Defaults to settings value.
        """
        settings = get_settings()
        self.chunk_size = chunk_size if chunk_size is not None else settings.chunk_size
        self.chunk_overlap = chunk_overlap if chunk_overlap is not None else settings.chunk_overlap
        
        if self.chunk_overlap >= self.chunk_size:
            raise ValueError("chunk_overlap must be less than chunk_size")
        
        logger.info(f"DocumentProcessor initialized with chunk_size={self.chunk_size}, overlap={self.chunk_overlap}")

    def process_file(self, file_path: str) -> ProcessedDocument:
        """
        Process a document file and return chunks.
        
        Args:
            file_path: Path to the document file.
        
        Returns:
            ProcessedDocument: Processed document with chunks.
        
        Raises:
            DocumentProcessingError: If processing fails.
        """
        path = Path(file_path)
        
        if not path.exists():
            raise DocumentProcessingError(
                f"File not found: {file_path}",
                f"File '{path.name}' does not exist.",
                recoverable=False
            )
        
        extension = path.suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise DocumentProcessingError(
                f"Unsupported file format: {extension}",
                f"Unsupported file format. Please use PDF, DOCX, TXT, or MD.",
                recoverable=False
            )
        
        source = sanitize_filename(path.name)
        
        try:
            if extension == '.pdf':
                text = self.process_pdf(file_path)
            elif extension == '.docx':
                text = self.process_docx(file_path)
            else:  # .txt or .md
                text = self.process_text(file_path)
        except DocumentProcessingError:
            raise
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            raise DocumentProcessingError(
                f"Failed to process {file_path}: {e}",
                f"Failed to process {path.name}: {str(e)}",
                recoverable=True
            )
        
        if not text or not text.strip():
            raise DocumentProcessingError(
                f"Empty document: {file_path}",
                f"Document '{path.name}' appears to be empty or unreadable.",
                recoverable=False
            )
        
        chunks = self.chunk_text(text, source)
        
        metadata = {
            "source": source,
            "original_path": str(path.absolute()),
            "extension": extension,
            "processed_at": datetime.now().isoformat(),
            "chunk_count": len(chunks)
        }
        
        logger.info(f"Processed {source}: {len(chunks)} chunks")
        return ProcessedDocument(source=source, chunks=chunks, metadata=metadata)
    
    def process_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file.
        
        Returns:
            str: Extracted text content.
        """
        try:
            from pypdf import PdfReader
        except ImportError:
            raise DocumentProcessingError(
                "pypdf not installed",
                "PDF processing requires pypdf. Please install it.",
                recoverable=False
            )
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
    
    def process_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file.
        
        Returns:
            str: Extracted text content preserving paragraph structure.
        """
        try:
            from docx import Document
        except ImportError:
            raise DocumentProcessingError(
                "python-docx not installed",
                "DOCX processing requires python-docx. Please install it.",
                recoverable=False
            )
        
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n\n".join(paragraphs)
    
    def process_text(self, file_path: str) -> str:
        """
        Read text from a TXT or MD file.
        
        Args:
            file_path: Path to the text file.
        
        Returns:
            str: Raw text content.
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def chunk_text(self, text: str, source: str) -> List[DocumentChunk]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: The text to chunk.
            source: Source filename for metadata.
        
        Returns:
            List[DocumentChunk]: List of document chunks.
        """
        if not text or not text.strip():
            return []
        
        chunks = []
        text = text.strip()
        
        # Calculate step size (chunk_size - overlap)
        step = self.chunk_size - self.chunk_overlap
        if step <= 0:
            step = 1
        
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = start + self.chunk_size
            chunk_content = text[start:end]
            
            # Only add non-empty chunks
            if chunk_content.strip():
                chunk = DocumentChunk(
                    id=generate_uuid(),
                    content=chunk_content,
                    metadata={
                        "source": source,
                        "chunk_index": chunk_index,
                        "created_at": datetime.now().isoformat()
                    },
                    embedding=None
                )
                chunks.append(chunk)
                chunk_index += 1
            
            start += step
            
            # Avoid infinite loop for very small texts
            if end >= len(text):
                break
        
        return chunks
